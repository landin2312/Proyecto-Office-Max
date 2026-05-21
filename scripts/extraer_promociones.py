from __future__ import annotations

import argparse
import gc
import hashlib
import logging
import shutil
import subprocess
import tempfile
from collections import Counter
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Any

import pandas as pd

_ORIGINAL_RMTREE = shutil.rmtree


def _rmtree_ignore_permission_errors(path: str, *args: Any, **kwargs: Any) -> None:
    try:
        _ORIGINAL_RMTREE(path, *args, **kwargs)
    except PermissionError:
        pass


shutil.rmtree = _rmtree_ignore_permission_errors

import pdfplumber
import regex as re
from docx import Document

try:
    import camelot
except Exception:  # pragma: no cover - camelot puede fallar si faltan binarios externos.
    camelot = None


COLUMNS = [
    "promo_id",
    "archivo_origen",
    "mes_carpeta",
    "fecha_comunicado",
    "fecha_inicio",
    "fecha_fin",
    "duracion_dias",
    "tiendas_participantes",
    "departamento",
    "mecanica",
    "comprador",
    "contacto_email",
    "promo_texto",
    "tipo_promo",
    "descuento_pct",
    "prod_nbr",
    "sku_sap",
    "prod_nm",
    "observaciones",
    "calidad_extraccion",
]

SUPPORTED_EXTENSIONS = {".pdf", ".doc", ".docx", ".xlsx"}
MONTHS = {
    "enero": 1,
    "febrero": 2,
    "marzo": 3,
    "abril": 4,
    "mayo": 5,
    "junio": 6,
    "julio": 7,
    "agosto": 8,
    "septiembre": 9,
    "setiembre": 9,
    "octubre": 10,
    "noviembre": 11,
    "diciembre": 12,
}


@dataclass
class ExtractedDocument:
    text: str
    tables: list[pd.DataFrame]
    warnings: list[str]


def null_if_empty(value: Any) -> Any:
    if value is None:
        return None
    if isinstance(value, float) and pd.isna(value):
        return None
    text = str(value).strip()
    return text if text else None


def compact_text(text: str | None, max_len: int = 3000) -> str | None:
    if not text:
        return None
    cleaned = re.sub(r"\s+", " ", text).strip()
    return cleaned[:max_len] if cleaned else None


def normalize_date(value: str | None, default_year: int | None = None) -> str | None:
    if not value:
        return None
    raw = re.sub(r"\s+", " ", value.strip().lower())
    raw = raw.replace(",", "")

    numeric = re.search(r"\b(\d{1,2})[./-](\d{1,2})[./-](\d{2,4})\b", raw)
    if numeric:
        day, month, year = numeric.groups()
        year_int = int(year)
        if year_int < 100:
            year_int += 2000
        try:
            return datetime(year_int, int(month), int(day)).date().isoformat()
        except ValueError:
            return None

    named = re.search(
        r"\b(\d{1,2})\s+(?:de\s+)?([a-záéíóúñ]+)\s+(?:de\s+)?(\d{4})\b",
        raw,
    )
    if named:
        day, month_name, year = named.groups()
        month = MONTHS.get(strip_accents(month_name))
        if month:
            try:
                return datetime(int(year), month, int(day)).date().isoformat()
            except ValueError:
                return None

    if default_year:
        short_named = re.search(r"\b(\d{1,2})\s+(?:de\s+)?([a-záéíóúñ]+)\b", raw)
        if short_named:
            day, month_name = short_named.groups()
            month = MONTHS.get(strip_accents(month_name))
            if month:
                try:
                    return datetime(default_year, month, int(day)).date().isoformat()
                except ValueError:
                    return None
    return None


def strip_accents(value: str) -> str:
    return (
        value.lower()
        .replace("á", "a")
        .replace("é", "e")
        .replace("í", "i")
        .replace("ó", "o")
        .replace("ú", "u")
        .replace("ü", "u")
    )


def extract_line_value(text: str, label_pattern: str) -> str | None:
    pattern = rf"(?im)^\s*{label_pattern}\s*[:\-]?\s*(.+?)\s*$"
    match = re.search(pattern, text)
    return compact_text(match.group(1), 500) if match else None


def extract_between_labels(text: str, label_pattern: str) -> str | None:
    labels = (
        r"NUMERO|NÚMERO|FECHA|VIGENCIA|TIENDAS\s+PARTICIPANTES|MEC[AÁ]NICA|"
        r"DEPARTAMENTO|COMPRADOR|CONTACTO|SKU|DESCUENTO"
    )
    pattern = rf"(?is)\b{label_pattern}\b\s*[:\-]?\s*(.+?)(?=\n\s*(?:{labels})\b\s*[:\-]?|\Z)"
    match = re.search(pattern, text)
    return compact_text(match.group(1), 1000) if match else None


def infer_promo_id(path: Path, text: str) -> str:
    number = extract_line_value(text, r"(?:NUMERO|N[UÚ]MERO)")
    if number:
        found = re.search(r"\d{3,}", number)
        if found:
            return found.group(0)
        return number
    prefix = re.match(r"^(\d{3,})", path.stem)
    if prefix:
        return prefix.group(1)
    digest = hashlib.sha1(str(path).encode("utf-8")).hexdigest()[:10]
    return f"sin_numero_{digest}"


def extract_dates(text: str, filename: str) -> tuple[str | None, str | None, str | None, int | None]:
    years = [int(y) for y in re.findall(r"\b(20\d{2})\b", text + " " + filename)]
    default_year = years[0] if years else None

    fecha_comunicado_raw = extract_line_value(text, r"FECHA")
    fecha_comunicado = normalize_date(fecha_comunicado_raw, default_year)

    vigencia = extract_between_labels(text, r"VIGENCIA") or extract_line_value(text, r"Vigencia")
    fecha_inicio = fecha_fin = None
    if vigencia:
        dates = re.findall(r"\d{1,2}[./-]\d{1,2}[./-]\d{2,4}", vigencia)
        if len(dates) >= 2:
            fecha_inicio = normalize_date(dates[0], default_year)
            fecha_fin = normalize_date(dates[1], default_year)
        else:
            range_named = re.search(
                r"(?i)\b(?:del|de)\s+(\d{1,2})\s+(?:al|a)\s+(\d{1,2})\s+(?:de\s+)?([a-záéíóúñ]+)\s+(?:de\s+)?(20\d{2})?",
                vigencia,
            )
            if range_named:
                start_day, end_day, month_name, year = range_named.groups()
                year_int = int(year) if year else default_year
                month = MONTHS.get(strip_accents(month_name))
                if year_int and month:
                    fecha_inicio = normalize_date(f"{start_day}/{month}/{year_int}")
                    fecha_fin = normalize_date(f"{end_day}/{month}/{year_int}")
            else:
                named_dates = re.findall(
                    r"\d{1,2}\s+(?:de\s+)?[a-záéíóúñ]+\s+(?:de\s+)?20\d{2}",
                    vigencia,
                    flags=re.I,
                )
                if len(named_dates) >= 2:
                    fecha_inicio = normalize_date(named_dates[0], default_year)
                    fecha_fin = normalize_date(named_dates[1], default_year)
                elif len(named_dates) == 1:
                    fecha_inicio = normalize_date(named_dates[0], default_year)

    duracion = None
    if fecha_inicio and fecha_fin:
        try:
            start = datetime.fromisoformat(fecha_inicio)
            end = datetime.fromisoformat(fecha_fin)
            duracion = (end - start).days + 1
        except ValueError:
            duracion = None
    return fecha_comunicado, fecha_inicio, fecha_fin, duracion


def extract_discount(text: str) -> str | None:
    labeled = extract_line_value(text, r"descuento")
    if labeled:
        normalized = normalize_discount_value(labeled)
        if normalized:
            return normalized
    match = re.search(r"(?i)\b(?:hasta\s+)?(\d{1,3}(?:[.,]\d+)?)\s*[%_]\s*(?:de\s*)?(?:desc|descuento|off)?\b", text)
    if match:
        return normalize_discount_value(match.group(1))
    match = re.search(r"(?i)\b([23]\s*x\s*[12])\b", text)
    if match:
        return normalize_discount_value(match.group(1))
    return None


def normalize_discount_value(value: Any) -> str | None:
    text = null_if_empty(value)
    if not text:
        return None
    cleaned = str(text).strip()
    cleaned = re.sub(r"\s+", " ", cleaned)
    cleaned = cleaned.replace("_", "%")

    combo = re.fullmatch(r"(?i)\s*([23])\s*x\s*([12])\s*", cleaned)
    if combo:
        return f"{combo.group(1)}x{combo.group(2)}"

    numeric = re.fullmatch(r"\s*(?:hasta\s+)?(\d{1,3}(?:[.,]\d+)?)\s*%?\s*(?:de\s*)?(?:desc(?:uento)?|off)?\s*", cleaned, flags=re.I)
    if not numeric:
        return None
    number_text = numeric.group(1).replace(",", ".")
    try:
        number = float(number_text)
    except ValueError:
        return None
    if number < 0 or number > 100:
        return None
    return str(int(number)) if number.is_integer() else str(number)


def infer_tipo_promo(text: str) -> str | None:
    lowered = strip_accents(text)
    if re.search(r"\b2\s*x\s*1\b|\b2x1\b", lowered):
        return "2x1"
    if re.search(r"\b3\s*x\s*2\b|\b3x2\b", lowered):
        return "3x2"
    if "%" in text or "descuento" in lowered or re.search(r"\boff\b", lowered):
        return "descuento"
    if "regalo" in lowered or "gratis" in lowered:
        return "regalo"
    if "cambio precio" in lowered or "actualizacion precio" in lowered:
        return "precio"
    return None


def extract_pdf(path: Path) -> ExtractedDocument:
    warnings: list[str] = []
    text_parts: list[str] = []
    tables: list[pd.DataFrame] = []

    with pdfplumber.open(path) as pdf:
        for page_number, page in enumerate(pdf.pages, start=1):
            text_parts.append(page.extract_text() or "")
            try:
                for table in page.extract_tables() or []:
                    if table:
                        tables.append(pd.DataFrame(table))
            except Exception as exc:
                warnings.append(f"pdfplumber tabla pagina {page_number}: {exc}")

    if camelot is not None:
        for flavor in ("lattice", "stream"):
            try:
                parsed = camelot.read_pdf(str(path), pages="all", flavor=flavor)
                for table in parsed:
                    if not table.df.empty:
                        tables.append(table.df)
                del parsed
                gc.collect()
            except Exception as exc:
                warnings.append(f"camelot {flavor}: {exc}")
    else:
        warnings.append("camelot no disponible")

    return ExtractedDocument("\n".join(text_parts), tables, warnings)


def extract_docx(path: Path) -> ExtractedDocument:
    doc = Document(path)
    text_parts = [p.text for p in doc.paragraphs if p.text]
    tables: list[pd.DataFrame] = []
    for table in doc.tables:
        rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
        if rows:
            tables.append(pd.DataFrame(rows))
            text_parts.extend(" | ".join(row) for row in rows)
    return ExtractedDocument("\n".join(text_parts), tables, [])


def extract_doc_with_textract(path: Path) -> ExtractedDocument:
    import textract

    raw = textract.process(str(path))
    text = raw.decode("utf-8", errors="replace") if isinstance(raw, bytes) else str(raw)
    return ExtractedDocument(text, [], ["doc extraido con textract"])


def extract_doc_with_antiword(path: Path) -> ExtractedDocument:
    completed = subprocess.run(
        ["antiword", str(path)],
        capture_output=True,
        check=True,
        text=True,
        encoding="utf-8",
        errors="replace",
    )
    return ExtractedDocument(completed.stdout, [], ["doc extraido con antiword"])


def word_com_app() -> Any:
    import win32com.client

    word = win32com.client.DispatchEx("Word.Application")
    word.Visible = False
    word.DisplayAlerts = 0
    return word


def find_soffice() -> str | None:
    found = shutil.which("soffice") or shutil.which("libreoffice")
    if found:
        return found
    candidates = [
        Path("C:/Program Files/LibreOffice/program/soffice.exe"),
        Path("C:/Program Files (x86)/LibreOffice/program/soffice.exe"),
        Path("C:/Program Files/OpenOffice 4/program/soffice.exe"),
        Path("C:/Program Files (x86)/OpenOffice 4/program/soffice.exe"),
    ]
    for candidate in candidates:
        if candidate.exists():
            return str(candidate)
    return None


def convert_doc_with_soffice_then_extract(path: Path) -> ExtractedDocument:
    soffice = find_soffice()
    if not soffice:
        raise FileNotFoundError("soffice/libreoffice no encontrado")

    tmp_dir = Path(tempfile.mkdtemp(prefix="doc_soffice_"))
    try:
        completed = subprocess.run(
            [
                soffice,
                "--headless",
                "--convert-to",
                "docx",
                "--outdir",
                str(tmp_dir.resolve()),
                str(path.resolve()),
            ],
            capture_output=True,
            text=True,
            encoding="utf-8",
            errors="replace",
            timeout=120,
        )
        if completed.returncode != 0:
            detail = (completed.stderr or completed.stdout or "").strip()
            raise RuntimeError(f"LibreOffice fallo con codigo {completed.returncode}: {detail}")
        converted = tmp_dir / f"{path.stem}.docx"
        if not converted.exists():
            matches = list(tmp_dir.glob("*.docx"))
            if not matches:
                raise FileNotFoundError(f"LibreOffice no genero DOCX en {tmp_dir}")
            converted = matches[0]
        extracted = extract_docx(converted)
        extracted.warnings.append("doc convertido temporalmente a docx con LibreOffice headless")
        return extracted
    finally:
        shutil.rmtree(tmp_dir, ignore_errors=True)


def extract_doc_with_pywin32(path: Path) -> ExtractedDocument:
    word = None
    doc = None
    try:
        word = word_com_app()
        doc = word.Documents.Open(str(path.resolve()), ReadOnly=True, ConfirmConversions=False)
        text = doc.Content.Text
        tables: list[pd.DataFrame] = []
        for table in doc.Tables:
            rows: list[list[str]] = []
            for row in table.Rows:
                values = []
                for cell in row.Cells:
                    values.append(str(cell.Range.Text).replace("\r", "").replace("\x07", "").strip())
                rows.append(values)
            if rows:
                tables.append(pd.DataFrame(rows))
        return ExtractedDocument(text, tables, ["doc extraido con pywin32"])
    finally:
        if doc is not None:
            doc.Close(False)
        if word is not None:
            word.Quit()


def convert_doc_to_docx_then_extract(path: Path) -> ExtractedDocument:
    word = None
    doc = None
    tmp_dir = Path(tempfile.mkdtemp(prefix="doc_convert_"))
    tmp_docx = tmp_dir / f"{path.stem}.docx"
    try:
        word = word_com_app()
        doc = word.Documents.Open(str(path.resolve()), ReadOnly=True, ConfirmConversions=False)
        doc.SaveAs2(str(tmp_docx.resolve()), FileFormat=16)
        doc.Close(False)
        doc = None
        extracted = extract_docx(tmp_docx)
        extracted.warnings.append("doc convertido temporalmente a docx con pywin32")
        return extracted
    finally:
        if doc is not None:
            doc.Close(False)
        if word is not None:
            word.Quit()
        shutil.rmtree(tmp_dir, ignore_errors=True)


def extract_doc(path: Path) -> ExtractedDocument:
    failures: list[str] = []
    methods = [
        ("libreoffice_headless_docx", convert_doc_with_soffice_then_extract),
        ("conversion_docx_temporal", convert_doc_to_docx_then_extract),
        ("textract", extract_doc_with_textract),
        ("antiword", extract_doc_with_antiword),
        ("pywin32", extract_doc_with_pywin32),
    ]
    for name, method in methods:
        try:
            extracted = method(path)
            if compact_text(extracted.text) or extracted.tables:
                extracted.warnings = failures + extracted.warnings
                return extracted
            failures.append(f"{name}: sin texto extraido")
        except Exception as exc:
            failures.append(f"{name}: {type(exc).__name__}: {exc}")
    raise RuntimeError("No se pudo extraer .doc con ningun metodo. " + " | ".join(failures))


def extract_xlsx(path: Path) -> ExtractedDocument:
    warnings: list[str] = []
    text_parts: list[str] = []
    tables: list[pd.DataFrame] = []
    workbook = pd.ExcelFile(path, engine="openpyxl")
    for sheet_name in workbook.sheet_names:
        try:
            sheet = pd.read_excel(workbook, sheet_name=sheet_name, header=None, dtype=str)
            sheet = sheet.dropna(how="all").dropna(axis=1, how="all")
            if sheet.empty:
                continue
            tables.append(sheet)
            text_parts.append(f"HOJA: {sheet_name}")
            for _, row in sheet.iterrows():
                values = [str(value).strip() for value in row.tolist() if null_if_empty(value)]
                if values:
                    text_parts.append(" | ".join(values))
        except Exception as exc:
            warnings.append(f"xlsx hoja {sheet_name}: {exc}")
    return ExtractedDocument("\n".join(text_parts), tables, warnings)


def normalize_table(df: pd.DataFrame) -> pd.DataFrame:
    cleaned = df.copy()
    cleaned = cleaned.map(lambda x: re.sub(r"\s+", " ", str(x)).strip() if x is not None else "")
    cleaned = cleaned.dropna(how="all")
    if cleaned.empty:
        return cleaned

    first_row = [strip_accents(str(x)) for x in cleaned.iloc[0].tolist()]
    if any(re.search(r"sku|sap|prod|producto|descripcion|descrip", cell) for cell in first_row):
        cleaned.columns = [strip_accents(str(x)).strip() or f"col_{i}" for i, x in enumerate(cleaned.iloc[0].tolist())]
        cleaned = cleaned.iloc[1:].reset_index(drop=True)
    else:
        cleaned.columns = [f"col_{i}" for i in range(len(cleaned.columns))]
    return cleaned


def pick_from_row(row: pd.Series, patterns: list[str]) -> str | None:
    for col, value in row.items():
        col_norm = strip_accents(str(col))
        if any(re.search(pattern, col_norm) for pattern in patterns):
            picked = null_if_empty(value)
            if picked:
                return picked
    return None


def pick_discount_from_row(row: pd.Series) -> str | None:
    allowed_patterns = [r"^%$", r"descuento", r"\bdesc\b", r"promo", r"promocion", r"mecanica"]
    blocked_patterns = [r"descripcion", r"descrip", r"producto", r"prod_nm", r"nombre", r"articulo"]
    for col, value in row.items():
        col_norm = strip_accents(str(col))
        if any(re.search(pattern, col_norm) for pattern in blocked_patterns):
            continue
        if any(re.search(pattern, col_norm) for pattern in allowed_patterns):
            normalized = normalize_discount_value(value)
            if normalized:
                return normalized
    for col, value in row.items():
        col_norm = strip_accents(str(col))
        if any(re.search(pattern, col_norm) for pattern in blocked_patterns):
            continue
        normalized = normalize_discount_value(value)
        if normalized:
            return normalized
    return None


def sku_rows_from_tables(tables: list[pd.DataFrame]) -> list[dict[str, str | None]]:
    rows: list[dict[str, str | None]] = []
    for raw in tables:
        table = normalize_table(raw)
        if table.empty:
            continue
        for _, row in table.iterrows():
            joined = " ".join(str(v) for v in row.tolist())
            row_context = strip_accents(" ".join([str(c) for c in table.columns]) + " " + joined)
            has_product_context = bool(re.search(r"\bsku\b|\bsap\b|codigo|producto|articulo|descrip|prod", row_context))
            if not has_product_context and not re.search(r"\b\d{7,10}\b", joined):
                continue
            sku = pick_from_row(row, [r"\bsku\b", r"sap", r"codigo", r"prod"])
            if not sku:
                sku = find_first_sku(joined, allow_generic=has_product_context)
            product = pick_from_row(row, [r"descripcion", r"producto", r"prod_nm", r"nombre"])
            if product and re.fullmatch(r"[-–—]+", product.strip()):
                product = None
            discount = pick_discount_from_row(row) or extract_discount(joined)
            if sku:
                rows.append(
                    {
                        "prod_nbr": sku,
                        "sku_sap": sku if re.search(r"\bsku\s*sap\b|\bsap\b", joined, re.I) else None,
                        "prod_nm": product,
                        "descuento_pct": discount,
                        "observaciones": compact_text(joined, 500),
                    }
                )
    return dedupe_rows(rows)


def find_first_sku(text: str, allow_generic: bool = False) -> str | None:
    labeled = re.search(r"(?i)\bSKU(?:\s*SAP)?\s*[:#\-]?\s*(\d{4,10})\b", text)
    if labeled:
        return labeled.group(1)
    if not allow_generic:
        long_generic = re.search(r"\b\d{7,10}\b", text)
        return long_generic.group(0) if long_generic else None
    generic = re.search(r"\b\d{5,10}\b", text)
    return generic.group(0) if generic else None


def sku_rows_from_text(text: str) -> list[dict[str, str | None]]:
    rows: list[dict[str, str | None]] = []
    for match in re.finditer(r"(?i)\bSKU(?:\s*SAP)?\s*[:#\-]?\s*(\d{4,10})([^\n]{0,180})", text):
        sku, tail = match.groups()
        rows.append(
            {
                "prod_nbr": sku,
                "sku_sap": sku if "sap" in strip_accents(match.group(0)) else None,
                "prod_nm": compact_text(tail, 180),
                "descuento_pct": extract_discount(match.group(0)),
                "observaciones": compact_text(match.group(0), 500),
            }
        )
    return dedupe_rows(rows)


def dedupe_rows(rows: list[dict[str, str | None]]) -> list[dict[str, str | None]]:
    seen: set[tuple[str | None, str | None]] = set()
    unique: list[dict[str, str | None]] = []
    for row in rows:
        key = (row.get("prod_nbr"), row.get("prod_nm"))
        if key not in seen:
            seen.add(key)
            unique.append(row)
    return unique


def build_base_record(path: Path, doc: ExtractedDocument) -> dict[str, Any]:
    text = doc.text
    fecha_comunicado, fecha_inicio, fecha_fin, duracion = extract_dates(text, path.name)
    mecanica = extract_between_labels(text, r"MEC[AÁ]NICA") or extract_line_value(text, r"MEC[AÁ]NICA")
    descuento = extract_discount(text)
    if not mecanica and descuento:
        mecanica = f"{descuento}% descuento"
    return {
        "promo_id": infer_promo_id(path, text),
        "archivo_origen": str(path),
        "mes_carpeta": path.parent.name,
        "fecha_comunicado": fecha_comunicado,
        "fecha_inicio": fecha_inicio,
        "fecha_fin": fecha_fin,
        "duracion_dias": duracion,
        "tiendas_participantes": extract_between_labels(text, r"Tiendas\s+participantes")
        or extract_line_value(text, r"Tiendas\s+participantes"),
        "departamento": extract_line_value(text, r"Departamento"),
        "mecanica": mecanica,
        "comprador": extract_line_value(text, r"Comprador"),
        "contacto_email": (re.search(r"[\w.+-]+@[\w.-]+\.[A-Za-z]{2,}", text) or [None])[0],
        "promo_texto": compact_text(text),
        "tipo_promo": infer_tipo_promo(text),
        "descuento_pct": descuento,
    }


def extraction_quality(base: dict[str, Any], sku_row: dict[str, Any], warnings: list[str]) -> str:
    score = 0
    for field in ("promo_id", "fecha_inicio", "fecha_fin", "mecanica", "tipo_promo"):
        if base.get(field):
            score += 1
    if sku_row.get("prod_nbr") or sku_row.get("sku_sap"):
        score += 2
    if sku_row.get("prod_nm"):
        score += 1
    if warnings:
        score -= 1
    if score >= 6:
        return "alta"
    if score >= 3:
        return "media"
    return "baja"


def process_file(path: Path) -> tuple[list[dict[str, Any]], list[str]]:
    if path.suffix.lower() == ".pdf":
        doc = extract_pdf(path)
    elif path.suffix.lower() == ".doc":
        doc = extract_doc(path)
    elif path.suffix.lower() == ".docx":
        doc = extract_docx(path)
    elif path.suffix.lower() == ".xlsx":
        doc = extract_xlsx(path)
    else:
        raise ValueError(f"extension no soportada: {path.suffix}")

    base = build_base_record(path, doc)
    sku_rows = sku_rows_from_tables(doc.tables)
    if not sku_rows:
        sku_rows = sku_rows_from_text(doc.text)
    if not sku_rows:
        return [], doc.warnings + ["sin SKU detectable"]

    records: list[dict[str, Any]] = []
    for sku_row in sku_rows:
        record = {col: None for col in COLUMNS}
        record.update(base)
        record.update({k: v for k, v in sku_row.items() if v})
        if not record.get("descuento_pct"):
            record["descuento_pct"] = base.get("descuento_pct")
        record["descuento_pct"] = normalize_discount_value(record.get("descuento_pct"))
        record["calidad_extraccion"] = extraction_quality(base, sku_row, doc.warnings)
        records.append(record)
    return records, doc.warnings


def discover_files(promociones_dir: Path, include_unsupported: bool = False) -> list[Path]:
    files = [p for p in promociones_dir.rglob("*") if p.is_file()]
    if include_unsupported:
        return sorted(files, key=lambda p: (p.parent.name, p.name.lower()))
    return sorted(
        [p for p in files if p.suffix.lower() in SUPPORTED_EXTENSIONS],
        key=lambda p: (p.parent.name, p.name.lower()),
    )


def extension_category(path: Path) -> str:
    ext = path.suffix.lower().lstrip(".")
    return ext if ext in {"pdf", "doc", "docx", "xlsx"} else "otros"


def category_counts(files: list[Path]) -> dict[str, int]:
    counts = Counter(extension_category(path) for path in files)
    return {category: counts.get(category, 0) for category in ("pdf", "doc", "docx", "xlsx", "otros")}


def write_outputs(records: list[dict[str, Any]], errors: list[dict[str, Any]], output_dir: Path) -> None:
    output_dir.mkdir(parents=True, exist_ok=True)
    master = pd.DataFrame(records, columns=COLUMNS)
    master.to_csv(output_dir / "promociones_master.csv", index=False, encoding="utf-8-sig")
    master.to_excel(output_dir / "promociones_master.xlsx", index=False)

    error_cols = ["archivo_origen", "mes_carpeta", "extension", "tipo_error", "detalle", "requiere_revision_manual"]
    pd.DataFrame(errors, columns=error_cols).to_csv(
        output_dir / "errores_extraccion.csv",
        index=False,
        encoding="utf-8-sig",
    )


def main() -> int:
    parser = argparse.ArgumentParser(description="Extrae promociones Office Max desde PDF, DOC, DOCX y XLSX.")
    parser.add_argument("--promociones-dir", default="Promociones", type=Path)
    parser.add_argument("--output-dir", default="output", type=Path)
    parser.add_argument("--limit", type=int, default=5, help="Numero maximo de archivos soportados a procesar.")
    parser.add_argument("--all", action="store_true", help="Procesa todos los PDF/DOC/DOCX/XLSX detectados.")
    parser.add_argument("--log-level", default="INFO")
    args = parser.parse_args()

    logging.basicConfig(level=getattr(logging, args.log_level.upper(), logging.INFO), format="%(levelname)s: %(message)s")
    all_files = discover_files(args.promociones_dir, include_unsupported=True)
    files = discover_files(args.promociones_dir)
    selected = files if args.all else files[: args.limit]
    unsupported = [
        p
        for p in args.promociones_dir.rglob("*")
        if p.is_file() and p.suffix.lower() not in SUPPORTED_EXTENSIONS
    ]
    temp_dir = args.output_dir / "_tmp"
    temp_dir.mkdir(parents=True, exist_ok=True)
    tempfile.tempdir = str(temp_dir.resolve())

    records: list[dict[str, Any]] = []
    errors: list[dict[str, Any]] = []
    processed_by_ext = Counter()
    for path in selected:
        logging.info("Procesando %s", path)
        processed_by_ext[extension_category(path)] += 1
        try:
            file_records, warnings = process_file(path)
            records.extend(file_records)
            if warnings:
                errors.append(
                    {
                        "archivo_origen": str(path),
                        "mes_carpeta": path.parent.name,
                        "extension": path.suffix.lower(),
                        "tipo_error": "advertencia",
                        "detalle": " | ".join(warnings),
                        "requiere_revision_manual": "SI" if not file_records else "NO",
                    }
                )
        except Exception as exc:
            errors.append(
                {
                    "archivo_origen": str(path),
                    "mes_carpeta": path.parent.name,
                    "extension": path.suffix.lower(),
                    "tipo_error": type(exc).__name__,
                    "detalle": str(exc),
                    "requiere_revision_manual": "SI",
                }
            )

    write_outputs(records, errors, args.output_dir)

    manual = sum(1 for error in errors if error.get("requiere_revision_manual") == "SI")
    failed = sum(1 for error in errors if error.get("tipo_error") not in {"advertencia", "extension_no_soportada"})
    detected_counts = category_counts(all_files)
    processed_counts = {category: processed_by_ext.get(category, 0) for category in ("pdf", "doc", "docx", "xlsx", "otros")}
    print("Resumen de extraccion")
    print(f"archivos detectados PDF/DOC/DOCX/XLSX: {len(files)}")
    print(f"archivos detectados no soportados: {len(unsupported)}")
    print(f"archivos procesados: {len(selected)}")
    print(f"filas extraidas: {len(records)}")
    print(f"archivos con errores: {failed}")
    print(f"archivos para revision manual: {manual}")
    print("conteo por extension detectada:")
    for category, count in detected_counts.items():
        print(f"  {category}: {count}")
    print("conteo por extension procesada:")
    for category, count in processed_counts.items():
        print(f"  {category}: {count}")
    print(f"salida master CSV: {args.output_dir / 'promociones_master.csv'}")
    print(f"salida master XLSX: {args.output_dir / 'promociones_master.xlsx'}")
    print(f"salida errores: {args.output_dir / 'errores_extraccion.csv'}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
